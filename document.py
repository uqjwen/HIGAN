from docx import Document
from docx.shared import Inches
from docx.shared import Pt
document = Document()
document.add_heading('Document Title', 0)  #插入标题
p = document.add_paragraph('A plain paragraph having some ')   #插入段落
run = p.add_run('bold')
# run.bold = True
run.font.size = Pt(20)
p.add_run(' and some ')
p.add_run('italic.').italic = True

p = document.add_paragraph('wenjh')
document.save('demo.docx')