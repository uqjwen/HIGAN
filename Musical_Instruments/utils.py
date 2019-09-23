import re
from process import readfile
import numpy as np 
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
def clean_str(string):
	"""
	Tokenization/string cleaning for all datasets except for SST.
	Original taken from https://github.com/yoonkim/CNN_sentence/blob/master/process_data.py
	"""

	string = re.sub(r"[^A-Za-z0-9(),!?\'\`]", " ", string)

	string = re.sub(r" \'", " ", string)
	string = re.sub(r"\' ", " ",string)

	string = re.sub(r"\'s", " s", string)
	string = re.sub(r"\'ve", " ve", string)
	# string = re.sub(r"n\'t", " n\'t", string)
	string = re.sub(r"\'re", " re", string)
	string = re.sub(r"\'d", " d", string)
	string = re.sub(r"\'ll", " ll", string)
	string = re.sub(r",", " , ", string)
	string = re.sub(r"!", " ! ", string)
	string = re.sub(r"\(", " \( ", string)
	string = re.sub(r"\)", " \) ", string)
	string = re.sub(r"\?", " \? ", string)
	string = re.sub(r"/", " / ", string)
	string = re.sub(r":", " :", string)
	string = re.sub(r"\s{2,}", " ", string)
	return string.strip()
def visual(data, uit, data_loader,utexts, itexts, u_texts, i_texts, filename):
	#idx: index of the sampled data point in the data_loader.train_uit
	idx2word = {v[1]:v[0] for v in data_loader.word2idx.items()}

	raw_data = readfile('./data/'+filename)


	document = Document()

	for idx, (doc_idx, doc) in enumerate(zip(utexts, u_texts)):
		word_atts = data[0][idx]
		raw_doc = raw_data[doc_idx-1]

		visual_single_doc(word_atts, doc, doc_idx, idx2word, raw_doc, document)

	document.add_paragraph('-----------------------------------------------------')
	for idx, (doc_idx, doc) in enumerate(zip(itexts, i_texts)):
		word_atts = data[1][idx]
		raw_doc = raw_data[doc_idx-1]

		visual_single_doc(word_atts, doc, doc_idx, idx2word, raw_doc, document)

	document.save('_'.join(map(str,uit))+'_atts.docx')
	doc_atts = []
	for user_layer in data[2]:
		doc_att = np.squeeze(user_layer)
		doc_atts.append(doc_att)
	for item_layer in data[3]:
		doc_att = np.squeeze(item_layer)
		doc_atts.append(doc_att)
	np.savetxt('_'.join(map(str,uit))+'_doc_atts.txt', doc_atts, fmt='%.5f')
		# doc_atts = [round(value,5) for value in doc_atts]
		# print(' '.join(map(str,doc_atts)))

	#data[0]: user word-level attention [1,6,60]
	#data[1]: item word-level attention 
	#data[3]: user document-level attention list of [1,6,1]
	#data[4]: item document-level attention 
	pass
def visual_single_doc(attentions, word_vec, doc_idx, idx2word, raw_doc, document):
	print('text id: ', doc_idx)
	import nltk
	raw_tokens = nltk.word_tokenize(clean_str(raw_doc['reviewText'].lower()))
	word2att = {}
	res = []
	for att, word_idx in zip(attentions, word_vec):
		if word_idx!=0:
			word = idx2word[word_idx]
			att = round(att,5)
			word2att[word] = max(word2att.get(word,0),att)

	atts = []
	for token in raw_tokens:
		att = word2att.get(token,0)	
		# res.append(token+':'+str(att))
		atts.append(att)
	single_document(raw_tokens, atts, document)
	# print(' '.join(res))


def single_document(tokens, atts, document):
	atts = np.array(atts)

	nz_atts = atts[atts!=0]
	min_size = 12
	max_size = 20
	if len(nz_atts)<2:
		return 

	min_att = min(nz_atts)
	max_att = max(nz_atts)

	if min_att == max_att:
		new_atts = np.array([15]*len(nz_atts))
	else:
		new_atts = 12+(max_size - min_size)/(max_att - min_att)*(nz_atts - min_att)

	atts[atts!=0] = new_atts
	atts = atts.astype(np.int32)

	p = document.add_paragraph('')

	for token,att in zip(tokens, atts):
		run = p.add_run(token+' ')
		if att>0:
			run.bold = True
			run.font.size = Pt(att)

